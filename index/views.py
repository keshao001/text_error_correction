from django.core.paginator import Paginator
from django.http import HttpResponseRedirect, JsonResponse
from django.shortcuts import render

from llm.TextHighlighter import TextHighlighter
from llm.qwen import ChatCompletion
from user.models import User
from .models import *
import os

workdir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
import datetime
import docx


def day_get():
    d = datetime.datetime.now()
    # 通过for 循环得到天数，如果想得到两周的时间，只需要把8改成15就可以了。
    for i in range(1, 8):
        oneday = datetime.timedelta(days=i)
        day = d - oneday
        date_to = datetime.datetime(day.year, day.month, day.day)
        yield str(date_to)[0:10]


def get_recent_seven_day():
    """
    获取最近的七天日期
    @return:
    """
    qq = day_get()
    day_list = []
    for obj in qq:
        day_list.append(obj)
    list_week_day = day_list[::-1]
    return list_week_day


def login(req):
    """
    跳转登录页面
    :param req:
    :return:
    """
    return render(req, 'login.html')


def register(req):
    """
    跳转注册页面
    :param req:
    :return:
    """
    return render(req, 'register.html')


def index(req):
    """
    跳转首页，并且返回用户最近7天的查询数据
    :param req:
    :return:
    """
    username = req.session['username']
    total_user = len(User.objects.all())
    date = datetime.datetime.today()
    month = date.month
    year = date.year
    seven_days = get_recent_seven_day()
    seven_data_dict = dict.fromkeys(seven_days, 0)
    for i in seven_days:
        d = datetime.timedelta
        r = Text.objects.filter(create_time__year=i.split('-')[0], create_time__month=i.split('-')[1],
                                create_time__day=i.split('-')[2]).all()
        if r:
            seven_data_dict[i] = r.count()
    seven_count_list = [seven_data_dict[x] for x in seven_days]
    return render(req, 'index.html', locals())


def login_out(req):
    """
    退出登录，返回首页
    :param req:
    :return:
    """
    del req.session['username']
    del req.session['role']
    del req.session['user_id']
    return HttpResponseRedirect('/')


def personal(req):
    username = req.session['username']
    role_id = req.session['role']
    user = User.objects.filter(name=username).first()

    return render(req, 'personal.html', locals())


def get_wb(request):
    """
    获取列表信息 | 模糊查询
    :param request:
    :return:
    """
    username = request.session.get('username')
    if not username:
        return JsonResponse({'code': 1, 'msg': '用户未登录', 'data': []})

    keyword = request.GET.get('name')
    page = request.GET.get("page", '')
    limit = request.GET.get("limit", '')
    role_id = request.GET.get('position','')

    response_data = {}
    response_data['code'] = 0
    response_data['msg'] = ''
    data = []
    if keyword is None:
        results_obj = Text.objects.all()
    else:
        results_obj = Text.objects.filter(owner=username,src__contains=keyword).all()
    paginator = Paginator(results_obj, limit)
    results = paginator.page(page)
    if results:
        for res in results:
            record = {
                "id": res.id, # 文本id
                "src": res.src, # 原文
                "dest": res.dest, # 更正后的文本
                "status": res.status, # 状态
                "owner": res.owner, # 用户
                'create_time': res.create_time.strftime('%Y-%m-%d %H:%m:%S'), # 创建时间
            }
            data.append(record)
        response_data['count'] = len(results_obj)
        response_data['data'] = data
    return JsonResponse(response_data)


def data(request):
    """
    跳转页面
    """
    username = request.session['username']
    role = int(request.session['role'])
    user_id = request.session['user_id']
    return render(request, 'data.html', locals())


def edit_data(request):
    """
    修改信息
    """
    response_data = {}
    user_id = request.POST.get('id')
    username = request.POST.get('username')
    phone = request.POST.get('phone')
    User.objects.filter(id=user_id).update(
        name=username,
        phone=phone)
    response_data['msg'] = 'success'
    return JsonResponse(response_data, status=201)


def del_wb(request):
    """
    删除文本
    """
    text_id = request.POST.get('id')
    result = Text.objects.filter(id=text_id).first()
    try:
        if not result:
            response_data = {'error': '删除失败！', 'message': '找不到id为%s' % text_id}
            return JsonResponse(response_data, status=403)
        result.delete()
        response_data = {'message': '删除成功！'}
        return JsonResponse(response_data, status=201)
    except Exception as e:
        response_data = {'message': '删除失败！'}
        return JsonResponse(response_data, status=403)


def del_doc(request):
    """
    删除文档
    """
    doc_id = request.POST.get('id')
    result = Document.objects.filter(id=doc_id).first()
    try:
        if not result:
            response_data = {'error': '删除失败！', 'message': '找不到id为%s' % doc_id}
            return JsonResponse(response_data, status=403)
        result.delete()
        response_data = {'message': '删除成功！'}
        return JsonResponse(response_data, status=201)
    except Exception as e:
        response_data = {'message': '删除失败！'}
        return JsonResponse(response_data, status=403)


def wdjc(request):
    username = request.session['username']
    role = int(request.session['role'])
    user_id = request.session['user_id']
    return render(request, 'wdjc.html', locals())


def wbjc(request):
    username = request.session['username']
    role = int(request.session['role'])
    user_id = request.session['user_id']
    return render(request, 'wbjc.html', locals())


def wbgl(request):
    username = request.session['username']
    role = int(request.session['role'])
    user_id = request.session['user_id']
    return render(request, 'wbgl.html', locals())

def save_file(file):
    if file is not None:
        file_name = os.path.join(workdir, 'static', 'uploadImg', file.name)
        with open(file_name, 'wb')as f:
            # chunks()每次读取数据默认 64k
            for chunk in file.chunks():
                f.write(chunk)
            f.close()
        return file_name
    else:
        return None

def wdgl(request):
    username = request.session['username']
    role = int(request.session['role'])
    user_id = request.session['user_id']
    return render(request, 'wdgl.html', locals())

def get_wdv1(request):
    keyword = request.GET.get('name')
    page = request.GET.get("page", '')
    limit = request.GET.get("limit", '')
    role_id = request.GET.get('position', '')

    response_data = {}
    response_data['code'] = 0
    response_data['msg'] = ''
    data = []

    if keyword is None:
        results_obj = Document.objects.all()
    else:
        results_obj = Document.objects.filter(name__contains=keyword).all()

    paginator = Paginator(results_obj, limit)
    results = paginator.page(page)

    if results:
        for res in results:
            record = {
                "id": res.id,
                "name": res.name,
                "src": res.src,
                "dest": res.dest,
                "owner": res.owner,
                "status": res.status,
                'create_time': res.create_time.strftime('%Y-%m-%d %H:%m:%S'),
            }
            data.append(record)

        response_data['count'] = len(results_obj)
        response_data['data'] = data
    return JsonResponse(response_data)

def get_wdv2(request):
    username = request.session['username']
    page = request.GET.get("page", '')
    limit = request.GET.get("pagesize", '')
    role_id = request.GET.get('position', '')

    response_data = {}
    response_data['code'] = 0
    response_data['msg'] = ''
    data = []

    results_obj = Document.objects.filter(owner=username).all()

    paginator = Paginator(results_obj, limit)
    results = paginator.page(page)

    if results:
        for res in results:
            record = {
                "id": res.id,
                "name": res.name,
                "src": res.src,
                "dest": res.dest,
                "owner": res.owner,
                "status": res.status,
                'create_time': res.create_time.strftime('%Y-%m-%d %H:%m:%S'),
            }
            data.append(record)

        response_data['count'] = len(results_obj)
        response_data['data'] = data
    return JsonResponse(response_data)
def correct_doc(request):
    if request.method == 'POST':
        doc = request.FILES.get('document')
    # text = 这理风景绣丽，而且天汽不错，我的心情各外舒畅!
    doc_content = docx.Document(doc)

    text = ""
    for paragraph in doc_content.paragraphs:
        text += paragraph.text + "\n"

    chat = ChatCompletion()
    result = chat.get_response(text)

    highlighter = TextHighlighter(text, result)
    highlighted_text = highlighter.highlight_differences()

    if text == result:
        status='无错误'
    else:
        status='有错误'

    Document.objects.create(name=doc.name,
                       src=text,
                       dest=result,
                       status=status,
                       owner=request.session.get('username', 'admin'),
                       )
    return JsonResponse({"result": highlighted_text, 'status': status, 'error': 0})

def correct_textv1(request):
    text = request.POST.get('text')
    #text = 这理风景绣丽，而且天汽不错，我的心情各外舒畅!

    chat = ChatCompletion()
    result = chat.get_response(text)

    highlighter = TextHighlighter(text, result)
    highlighted_text = highlighter.highlight_differences()

    if text ==result:
        status = '无错误'
    else:
        status = '有错误'

    Text.objects.create(
                       src=text,
                       dest=result,
                       status=status,
                       owner=request.session.get('username', 'admin'),
                       )

    return JsonResponse({"result": highlighted_text, 'status': status, 'error': 0})

# def correct_textv2(request):
#     text = request.POST.get('text')
#     #text = 涡伦风扇发动机通过压起机和燃烧室产生推力，推动飞机前进并提高效率，大大提高比腿力。
#
#     # RAG增强
#     # 用于增强的文本
#     pdf_path = "C:\\Users\\18817\\Desktop\\text_error_correction\\llm\\RAG\\RAGResources\\RAG.pdf"
#     # 分词模型
#     model_path = "C:\\Users\\18817\\Desktop\\text_error_correction\\llm\\RAG\\EmbeddingModels\\m3e-base"
#     # 创建 RAGPromptEnhancer 实例
#     enhancer = RAGPromptEnhancer(pdf_path, model_path)
#     # 增强的文本
#     enhanced_text = enhancer.enhance_prompt(text)
#
#     # 大模型更正
#     chat = ChatCompletion()
#     result = chat.get_response(enhanced_text)
#
#     highlighter = TextHighlighter(text, result)
#     highlighted_text = highlighter.highlight_differences()
#
#     if text ==result:
#         status = '无错误'
#     else:
#         status = '有错误'
#
#     Text.objects.create(
#                        src=text,
#                        dest=result,
#                        status=status,
#                        owner=request.session.get('username', 'admin'),
#                        )
#
#     return JsonResponse({"result": highlighted_text, 'status': status, 'error': 0})

def getdoccorrectresult(request,doc_id):
    doc = Document.objects.filter(id=doc_id).first()
    result = doc.dest
    return JsonResponse({"result":result})

def get_knowledge(request):
    username = request.session['username']
    response_data = {}
    response_data['code'] = 0
    response_data['msg'] = ''
    data = []

    results = Document.objects.filter(owner=username).all()

    if results:
        for res in results:
            if res.status == "知识库":
                record = {
                    "id": res.id,
                    "name": res.name,
                    "src": res.src,
                    "dest": res.dest,
                    "owner": res.owner,
                    "status": res.status,
                    'create_time': res.create_time.strftime('%Y-%m-%d %H:%m:%S'),
                }
                data.append(record)
        response_data['count'] = len(results)
        response_data['data'] = data
    return JsonResponse(response_data)

def upload_knowledge(request):
    if request.method == 'POST':
        doc = request.FILES.get('document')
    doc_content = docx.Document(doc)
    status = '知识库'
    Document.objects.create(name=doc.name,
                       src=doc_content,
                       dest="",
                       status=status,
                       owner=request.session.get('username', 'admin'),
                       )
    return JsonResponse({'msg': '上传成功'})